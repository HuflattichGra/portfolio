#!/usr/bin/env python3
"""
Extract text content from DOCX resume file
"""

from docx import Document
import sys

def extract_text_from_docx(file_path):
    """Extract all text from a DOCX file"""
    try:
        doc = Document(file_path)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Also extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text.strip())
        
        return '\n'.join(text_content)
    
    except Exception as e:
        return f"Error reading DOCX file: {str(e)}"

if __name__ == "__main__":
    docx_file = "Junran_CV_ra.docx"
    content = extract_text_from_docx(docx_file)
    print(content)
